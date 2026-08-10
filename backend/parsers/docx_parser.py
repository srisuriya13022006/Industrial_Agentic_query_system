from docx import Document


class DocxParser:
    """
    Extracts text content from DOCX files using python-docx.
    Handles paragraphs and tables.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract all text from a DOCX file.
        Includes paragraph text and table content.
        """

        doc = Document(file_path)

        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:

            content = paragraph.text.strip()

            if content:
                text_parts.append(content)

        # Extract tables
        for table in doc.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
